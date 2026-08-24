import sys
if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os
import random
import openpyxl
from openpyxl.styles import Alignment, Font as ExcelFont, PatternFill
from openpyxl.utils import get_column_letter

def set_cell_background(cell, fill_color):
    """Sets background color of a cell (fill_color should be hex string like 'F2F2F2')"""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_font(run, font_name="Times New Roman", size_pt=11, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic

def add_paragraph_with_run(doc, text="", font_name="Times New Roman", size_pt=11, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=5, line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        set_font(run, font_name, size_pt, bold, italic)
    return p

def create_signature_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    widths = [Inches(1.5), Inches(2.75), Inches(2.75)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            
    cells = table.rows[0].cells
    
    p1 = cells[0].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run("Điểm bài thi:\n")
    set_font(r1, font_name="Times New Roman", bold=True, size_pt=11)
    for _ in range(3):
        p_empty = cells[0].add_paragraph()
        p_empty.paragraph_format.space_after = Pt(0)
        
    p2 = cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run("Chữ ký Cán bộ chấm thi:\n")
    set_font(r2, font_name="Times New Roman", bold=True, size_pt=10.5)
    r2_sub = p2.add_run("(Ký và ghi rõ họ tên)")
    set_font(r2_sub, font_name="Times New Roman", italic=True, size_pt=9.5)
    for _ in range(2):
        p_empty = cells[1].add_paragraph()
        p_empty.paragraph_format.space_after = Pt(0)
        
    p3 = cells[2].paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(2)
    r3 = p3.add_run("Chữ ký Cán bộ coi thi:\n")
    set_font(r3, font_name="Times New Roman", bold=True, size_pt=10.5)
    r3_sub = p3.add_run("(Ký và ghi rõ họ tên)")
    set_font(r3_sub, font_name="Times New Roman", italic=True, size_pt=9.5)
    for _ in range(2):
        p_empty = cells[2].add_paragraph()
        p_empty.paragraph_format.space_after = Pt(0)
        
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(6)

def render_header(doc, position, candidate_name, candidate_id, exam_date, dept=""):
    top_p1 = doc.add_paragraph()
    top_p1.paragraph_format.space_after = Pt(2)
    top_p1.paragraph_format.line_spacing = 1.15
    r_hosp = top_p1.add_run("BỆNH VIỆN ĐA KHOA THIỆN HẠNH\n")
    set_font(r_hosp, font_name="Times New Roman", bold=True, size_pt=11)
    r_council = top_p1.add_run("HỘI ĐỒNG THI TUYỂN DỤNG")
    set_font(r_council, font_name="Times New Roman", bold=True, size_pt=11)

    create_signature_table(doc)
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(3)
    r_title = title_p.add_run("BÀI THI KỸ NĂNG TIN HỌC / THAO TÁC PHẦN MỀM")
    set_font(r_title, font_name="Times New Roman", bold=True, size_pt=14.5)
    
    sub_title_p = doc.add_paragraph()
    sub_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_title_p.paragraph_format.space_after = Pt(10)
    dept_str = f" - Khoa / Phòng: {dept}" if dept else ""
    r_sub = sub_title_p.add_run(f"Vị trí dự thi: {position}{dept_str}")
    set_font(r_sub, font_name="Times New Roman", italic=True, size_pt=11.5)
    
    # Candidate Info Table Box (2 columns)
    info_table = doc.add_table(rows=2, cols=2)
    info_table.style = 'Table Grid'
    widths = [Inches(4.2), Inches(2.8)]
    for row in info_table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Row 1, Col 0: Họ tên
    c00 = info_table.rows[0].cells[0]
    p00 = c00.paragraphs[0]
    p00.paragraph_format.space_after = Pt(2)
    p00.paragraph_format.line_spacing = 1.15
    r_l1 = p00.add_run("Họ và tên thí sinh: ")
    set_font(r_l1, font_name="Times New Roman", size_pt=11)
    r_val1 = p00.add_run(f"{candidate_name or '……………………………………………………'}")
    set_font(r_val1, font_name="Times New Roman", bold=True, size_pt=11.5)

    # Row 1, Col 1: SBD
    c01 = info_table.rows[0].cells[1]
    p01 = c01.paragraphs[0]
    p01.paragraph_format.space_after = Pt(2)
    p01.paragraph_format.line_spacing = 1.15
    r_l2 = p01.add_run("Số báo danh (SBD): ")
    set_font(r_l2, font_name="Times New Roman", size_pt=11)
    r_val2 = p01.add_run(f"{candidate_id or '……………'}")
    set_font(r_val2, font_name="Times New Roman", bold=True, size_pt=11.5)

    # Row 2, Col 0: Khoa phòng dự tuyển
    c10 = info_table.rows[1].cells[0]
    p10 = c10.paragraphs[0]
    p10.paragraph_format.space_after = Pt(2)
    p10.paragraph_format.line_spacing = 1.15
    r_l3 = p10.add_run("Khoa / Phòng dự tuyển: ")
    set_font(r_l3, font_name="Times New Roman", size_pt=11)
    r_val3 = p10.add_run(f"{dept or 'Toàn viện'}")
    set_font(r_val3, font_name="Times New Roman", bold=True, size_pt=11)

    # Row 2, Col 1: Ngày thi
    c11 = info_table.rows[1].cells[1]
    p11 = c11.paragraphs[0]
    p11.paragraph_format.space_after = Pt(2)
    p11.paragraph_format.line_spacing = 1.15
    r_l4 = p11.add_run("Ngày thi: ")
    set_font(r_l4, font_name="Times New Roman", size_pt=11)
    r_val4 = p11.add_run(f"{exam_date}")
    set_font(r_val4, font_name="Times New Roman", size_pt=11)

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(8)

# 1. NURSE_INPATIENT DOCX
def generate_nurse_inpatient_docx(doc, data, candidate_name, candidate_id, position, scores):
    render_header(doc, position, candidate_name, candidate_id, data["exam_date"])
    q_num = 1
    
    if len(scores) > 0 and scores[0] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Nhận bệnh nhân vào khoa. ({scores[0]}đ):", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, f"   * {data['bn_vp']['TenBenhNhan']} (Viện phí; mã y tế do HIS sinh khi nạp dữ liệu)", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, f"   * {data['bn_bhyt']['TenBenhNhan']} (BHYT; mã y tế do HIS sinh khi nạp dữ liệu)", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "- SVV Bn VP       : ………………………………", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, "- SVV Bn BHYT : ……………………………….", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 1 and scores[1] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Chỉ định các dịch vụ CLS, thủ thuật cho bệnh nhân (Bệnh nhân BHYT): ({scores[1]}đ)", bold=True, size_pt=11, space_after=4)
        services_text = ", ".join([s["TenDichVu"] for s in data["imaging_services"]])
        if services_text:
            services_text += ". "
        services_text += "Các dịch vụ xét nghiệm: " + "; ".join([s["TenDichVu"] for s in data["lab_services"]])
        add_paragraph_with_run(doc, services_text, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 2 and scores[2] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Thực hiện các thao tác y lệnh (Bệnh nhân BHYT): ({scores[2]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "- Lên y lệnh Thuốc & VTYT đính kèm", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, "- Lập phiếu lĩnh dược, Nhập nội bộ, Xuất sử dụng cho bệnh nhân", size_pt=11, space_after=4)
        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        headers = ["STT", "Mã Dược", "Tên thuốc / Vật tư y tế", "ĐVT", "SL", "Nguồn"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.paragraphs[0].text = h
            set_cell_background(cell, "EAEAEA")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.runs[0], bold=True, size_pt=10)
            
        drugs_list = data["bhyt_drugs"] + data["vp_supplies"]
        for idx, item in enumerate(drugs_list):
            row_cells = table.add_row().cells
            row_cells[0].paragraphs[0].text = str(idx + 1)
            row_cells[1].paragraphs[0].text = item["MaDuoc"]
            row_cells[2].paragraphs[0].text = item["TenDuoc"]
            row_cells[3].paragraphs[0].text = item["DVTTinh"]
            row_cells[4].paragraphs[0].text = str(int(random.choice([1, 2, 3])))
            row_cells[5].paragraphs[0].text = item["Nguon"]
            for col_idx in [0, 1, 3, 4, 5]:
                row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 3 and scores[3] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Bệnh nhân trả thuốc và tổng hợp phiếu trả: Thuốc cần trả {data['returned_drug']['TenDuoc']} (Số lượng: 1 {data['returned_drug']['DVTTinh']}). ({scores[3]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 4 and scores[4] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Đổi dịch vụ {data['swap_out_service']['TenDichVu']} thành dịch vụ {data['swap_in_service']['TenDichVu']} và thêm dịch vụ {data['added_service']['TenDichVu']}. ({scores[4]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 5 and scores[5] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Báo số tồn kho dược từ {data['check_start_date']} đến {data['check_end_date']} : Tên dược: {data['check_drug']['TenDuoc']} - Tồn cuối: ……………. ({scores[5]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 6 and scores[6] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Chuyển bệnh nhân từ khoa {position} đến {data['transfer_target_dept']} (bệnh nhân nhập viện trước {data['check_end_date']}), Ghi lại SVV: …………………, Họ tên BN: ……………………………… ({scores[6]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 7 and scores[7] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Thực hiện cho bệnh nhân ra viện (bệnh nhân nhập viện trước {data['check_end_date']}) : Ghi lại SVV: ……………………… ({scores[7]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    add_paragraph_with_run(doc, f"Lưu ý : User đăng nhập chương trình: {data['user']['TenDangNhap']}; PassWord : {data['user']['MatKhau']}", bold=True, size_pt=11, space_after=12)

# 2. RECEPTIONIST_INPATIENT DOCX (Lễ tân Khoa Điều trị Nội trú)
def generate_receptionist_inpatient_docx(doc, data, candidate_name, candidate_id, position, scores):
    render_header(doc, position, candidate_name, candidate_id, data["exam_date"])
    q_num = 1

    if len(scores) > 0 and scores[0] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Nhận bệnh nhân vào khoa điều trị ({scores[0]}đ):", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, f"   * BN BHYT: {data['bn_bhyt']['TenBenhNhan']} (mã y tế do HIS sinh khi nạp dữ liệu)", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, f"   * BN Viện phí: {data['bn_vp']['TenBenhNhan']} (mã y tế do HIS sinh khi nạp dữ liệu)", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "- Số vào viện BN BHYT: ………………………… | Số vào viện BN VP: …………………………", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 1 and scores[1] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Thu tạm ứng cho các bệnh nhân tại khoa điều trị ({scores[1]}đ):", bold=True, size_pt=11, space_after=4)
        for case in data["deposit_cases"]:
            add_paragraph_with_run(doc, f"   * Bệnh nhân số BA/STN: {case['ma_ba']}, Số tiền tạm ứng: {case['tien']} đồng", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 2 and scores[2] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Tra cứu đợt khám/nằm viện gần nhất của bệnh nhân {data['search_patient']['TenBenhNhan']} (Năm sinh: {data['search_patient']['NgaySinh']}): ({scores[2]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "   Ghi lại: Khoa điều trị: ……………………, Phòng: …………, Số giường: …………", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    add_paragraph_with_run(doc, f"Lưu ý: Dùng User: {data['user']['TenDangNhap']}; PassWord: {data['user']['MatKhau']} để thực hiện.", bold=True, size_pt=11)

# 3. NURSE_DIRECT_RECEPTION DOCX (Điều dưỡng Cấp cứu, Sản, Nhi)
def generate_nurse_direct_reception_docx(doc, data, candidate_name, candidate_id, position, scores):
    render_header(doc, position, candidate_name, candidate_id, data["exam_date"])
    q_num = 1

    if len(scores) > 0 and scores[0] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Trực tiếp tiếp nhận mới 2 bệnh nhân vào khoa (nhập đầy đủ thông tin hành chính): ({scores[0]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, f"1. BHYT: {data['bn_bhyt']['TenBenhNhan']} - NS: {data['bn_bhyt']['NgaySinh']} - Địa chỉ: {data['bn_bhyt']['DiaChiLienHe']} - Số BHYT: {data['bn_bhyt']['SoBHYT']}", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, f"2. Viện phí: {data['bn_vp']['TenBenhNhan']} - NS: {data['bn_vp']['NgaySinh']} - Địa chỉ: {data['bn_vp']['DiaChiLienHe']}", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "- STN Bn BHYT: ……………………………… | STN Bn VP: ………………………………", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 1 and scores[1] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Chỉ định các dịch vụ CLS cấp cứu/chuyên khoa cho bệnh nhân Viện phí: ({scores[1]}đ)", bold=True, size_pt=11, space_after=4)
        cls_text = ", ".join([s["TenDichVu"] for s in data["img_services"]]) + "; " + ", ".join([s["TenDichVu"] for s in data["lab_services"]])
        add_paragraph_with_run(doc, f"   Dịch vụ chỉ định: {cls_text}", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 2 and scores[2] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Lên y lệnh Thuốc & VTYT, lập phiếu lĩnh và xuất sử dụng (Bệnh nhân BHYT): ({scores[2]}đ)", bold=True, size_pt=11, space_after=4)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        headers = ["STT", "Mã Dược", "Tên Thuốc / Vật tư", "ĐVT", "Số lượng"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.paragraphs[0].text = h
            set_cell_background(cell, "EAEAEA")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.runs[0], bold=True, size_pt=10)
        items = data["drugs"] + data["supplies"]
        for idx, item in enumerate(items):
            row_cells = table.add_row().cells
            row_cells[0].paragraphs[0].text = str(idx + 1)
            row_cells[1].paragraphs[0].text = item["MaDuoc"]
            row_cells[2].paragraphs[0].text = item["TenDuoc"]
            row_cells[3].paragraphs[0].text = item["DVTTinh"]
            row_cells[4].paragraphs[0].text = str(random.choice([1, 2, 3]))
            for col_idx in [0, 1, 3, 4]:
                row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 3 and scores[3] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Trả thuốc {data['returned_drug']['TenDuoc']} (Số lượng: 1) và lập phiếu trả cho BN BHYT. ({scores[3]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 4 and scores[4] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Báo cáo tồn kho cuối của thuốc {data['check_drug']['TenDuoc']} từ {data['check_start_date']} đến {data['check_end_date']}. ({scores[4]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 5 and scores[5] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Cho bệnh nhân nhập viện trước ngày {data['check_end_date']} ra viện: SVV BN: …………………… ({scores[5]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    add_paragraph_with_run(doc, f"Lưu ý: Dùng User: {data['user']['TenDangNhap']}; PassWord: {data['user']['MatKhau']} để thực hiện.", bold=True, size_pt=11)

# 4. RECEPTIONIST_DIRECT DOCX (Lễ tân Cấp cứu, Sản, Nhi, Khám bệnh)
def generate_receptionist_direct_docx(doc, data, candidate_name, candidate_id, position, scores):
    render_header(doc, position, candidate_name, candidate_id, data["exam_date"])
    q_num = 1

    if len(scores) > 0 and scores[0] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Trực tiếp tiếp nhận mới 2 bệnh nhân (nhập đầy đủ thông tin hành chính & thẻ BHYT) ({scores[0]}đ):", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, f"1. BHYT: {data['bn_bhyt']['TenBenhNhan']} - NS: {data['bn_bhyt']['NgaySinh']} - Địa chỉ: {data['bn_bhyt']['DiaChiLienHe']} - Thẻ BHYT: {data['bn_bhyt']['SoBHYT']}", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, f"2. Viện phí: {data['bn_vp']['TenBenhNhan']} - NS: {data['bn_vp']['NgaySinh']} - Địa chỉ: {data['bn_vp']['DiaChiLienHe']}", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Ghi lại: STN BN BHYT: …………………… | STN BN VP: ……………………", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 1 and scores[1] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Thu phí khám và tạo chỉ định CLS cho bệnh nhân Viện phí ({scores[1]}đ):", bold=True, size_pt=11, space_after=4)
        cls_text = ", ".join([s["TenDichVu"] for s in data["img_services"]]) + "; " + ", ".join([s["TenDichVu"] for s in data["lab_services"]])
        add_paragraph_with_run(doc, f"   Chỉ định: {cls_text}", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, "   Thực hiện thu phí cho các dịch vụ vừa chỉ định.", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 2 and scores[2] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Thu tạm ứng cho bệnh nhân ({scores[2]}đ):", bold=True, size_pt=11, space_after=4)
        for case in data["deposit_cases"]:
            add_paragraph_with_run(doc, f"   * Bệnh nhân số BA/STN: {case['ma_ba']}, Số tiền tạm ứng: {case['tien']} đồng", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 3 and scores[3] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Tra cứu đợt khám gần nhất của bệnh nhân {data['search_patient']['TenBenhNhan']} (Năm sinh: {data['search_patient']['NgaySinh']}): ({scores[3]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "   Ghi lại: Khoa điều trị: ……………………, Phòng: …………, Số giường: …………", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    add_paragraph_with_run(doc, f"Lưu ý: Dùng User: {data['user']['TenDangNhap']}; PassWord: {data['user']['MatKhau']} để thực hiện.", bold=True, size_pt=11)

# 5. NURSE_OUTPATIENT DOCX
def generate_nurse_outpatient_docx(doc, data, candidate_name, candidate_id, position, scores):
    render_header(doc, position, candidate_name, candidate_id, data["exam_date"])
    q_num = 1

    if len(scores) > 0 and scores[0] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Tiếp nhận bệnh nhân khám ngoại trú BHYT & Viện phí vào phòng khám: ({scores[0]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, f"1. BN BHYT: {data['bn_bhyt']['TenBenhNhan']} - Số BHYT: {data['bn_bhyt']['SoBHYT']}", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, f"2. BN VP: {data['bn_vp']['TenBenhNhan']} - Địa chỉ: {data['bn_vp']['DiaChiLienHe']}", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 1 and scores[1] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Chỉ định các dịch vụ cận lâm sàng phòng khám cho bệnh nhân Viện phí: ({scores[1]}đ)", bold=True, size_pt=11, space_after=4)
        svc_str = ", ".join([s["TenDichVu"] for s in data["services"]])
        add_paragraph_with_run(doc, f"   Dịch vụ: {svc_str}", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 2 and scores[2] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Kê đơn thuốc ngoại trú cho bệnh nhân BHYT theo danh sách sau: ({scores[2]}đ)", bold=True, size_pt=11, space_after=4)
        for idx, d in enumerate(data["drugs"]):
            add_paragraph_with_run(doc, f"   {idx+1}. {d['TenDuoc']} ({d['DVTTinh']}) - SL: {random.choice([10, 20, 30])}", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 3 and scores[3] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Thu phí khám & dịch vụ cận lâm sàng cho bệnh nhân Viện phí. ({scores[3]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 4 and scores[4] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Đặt lịch hẹn tái khám cho bệnh nhân BHYT sau 7 ngày. ({scores[4]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    add_paragraph_with_run(doc, f"Lưu ý: Dùng User: {data['user']['TenDangNhap']}; PassWord: {data['user']['MatKhau']} để thực hiện.", bold=True, size_pt=11)

# 6. TECHNICIAN DOCX
def generate_technician_docx(doc, data, candidate_name, candidate_id, position, scores):
    render_header(doc, position, candidate_name, candidate_id, data["exam_date"])
    q_num = 1

    if len(scores) > 0 and scores[0] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Nhập và trả kết quả cận lâm sàng cho 5 bệnh nhân theo mẫu kết quả sau ({scores[0]}đ):", bold=True, size_pt=11, space_after=4)
        
        for idx, item in enumerate(data["list_patients_services"]):
            add_paragraph_with_run(doc, f"BN {idx+1}: {item['patient']['TenBenhNhan']} - Mã Y Tế: {item['patient']['MaYTe']} | Yêu cầu: {item['service']['TenDichVu']}", bold=True, size_pt=10, space_after=2)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            headers = ["Tên chỉ số / Hạng mục", "Kết quả mẫu", "Đơn vị tính", "Giá trị tham chiếu"]
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.paragraphs[0].text = h
                set_cell_background(cell, "F0F0F0")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_font(p.runs[0], bold=True, size_pt=9)
                
            sample_results = [
                ("WBC (Bạch cầu)", "6.8", "G/L", "4.0 - 10.0"),
                ("RBC (Hồng cầu)", "4.5", "T/L", "3.8 - 5.4"),
                ("HGB (Huyết sắc tố)", "138", "g/L", "120 - 160"),
                ("PLT (Tiểu cầu)", "245", "G/L", "150 - 400"),
                ("Glucose máu", "5.4", "mmol/L", "3.9 - 6.4")
            ]
            for row_data in sample_results:
                row_cells = table.add_row().cells
                for col_i, val in enumerate(row_data):
                    row_cells[col_i].paragraphs[0].text = val
                    set_font(row_cells[col_i].paragraphs[0].runs[0], size_pt=9)
                    if col_i in [1, 2, 3]:
                        row_cells[col_i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 1 and scores[1] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Lên phiếu lĩnh hóa chất/VTYT tiêu hao, nhập kho nội bộ và xuất sử dụng: ({scores[1]}đ)", bold=True, size_pt=11, space_after=4)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        headers = ["STT", "Mã", "Tên hóa chất / VTYT", "ĐVT", "Số lượng"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.paragraphs[0].text = h
            set_cell_background(cell, "EAEAEA")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.runs[0], bold=True, size_pt=10)
            
        for idx, item in enumerate(data["supplies"]):
            row_cells = table.add_row().cells
            row_cells[0].paragraphs[0].text = str(idx + 1)
            row_cells[1].paragraphs[0].text = item["MaDuoc"]
            row_cells[2].paragraphs[0].text = item["TenDuoc"]
            row_cells[3].paragraphs[0].text = item["DVTTinh"]
            row_cells[4].paragraphs[0].text = str(random.choice([2, 5, 10]))
            for col_idx in [0, 1, 3, 4]:
                row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    add_paragraph_with_run(doc, f"Lưu ý: Dùng User: {data['user']['TenDangNhap']}; PassWord: {data['user']['MatKhau']} để thực hiện.", bold=True, size_pt=11)

# 7. CASHIER_COUNTER DOCX (Thu ngân Quầy Thu phí)
def generate_cashier_counter_docx(doc, data, candidate_name, candidate_id, position, scores):
    render_header(doc, position, candidate_name, candidate_id, data["exam_date"])
    q_num = 1

    if len(scores) > 0 and scores[0] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Tiếp nhận 2 bệnh nhân BHYT và Viện Phí tại Quầy Thu phí ({scores[0]}đ):", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, f"1. BHYT: {data['bn_bhyt']['TenBenhNhan']} - Số BHYT: {data['bn_bhyt']['SoBHYT']}", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, f"2. Viện phí: {data['bn_vp']['TenBenhNhan']} - Địa chỉ: {data['bn_vp']['DiaChiLienHe']}", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Ghi lại: STN BN BHYT: …………………… | STN BN VP: ……………………", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 1 and scores[1] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Thu phí khám và tạo chỉ định CLS cho bệnh nhân Viện phí ({scores[1]}đ):", bold=True, size_pt=11, space_after=4)
        cls_text = ", ".join([s["TenDichVu"] for s in data["img_services"]]) + "; " + ", ".join([s["TenDichVu"] for s in data["lab_services"]])
        add_paragraph_with_run(doc, f"   Chỉ định: {cls_text}", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, "   Thực hiện thu phí cho các dịch vụ vừa chỉ định.", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 2 and scores[2] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Thu tạm ứng cho bệnh nhân ({scores[2]}đ):", bold=True, size_pt=11, space_after=4)
        for case in data["deposit_cases"]:
            add_paragraph_with_run(doc, f"   * Bệnh nhân số BA/STN: {case['ma_ba']}, Số tiền tạm ứng: {case['tien']} đồng", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 3 and scores[3] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Tra cứu đợt khám/nằm viện gần nhất của bệnh nhân {data['search_patient']['TenBenhNhan']} (Năm sinh: {data['search_patient']['NgaySinh']}): ({scores[3]}đ)", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "   Ghi lại: Khoa điều trị: ……………………, Phòng: …………, Số giường: …………", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    add_paragraph_with_run(doc, f"Lưu ý: Dùng User: {data['user']['TenDangNhap']}; PassWord: {data['user']['MatKhau']} để thực hiện.", bold=True, size_pt=11)

# 8. OFFICE_ADMIN DOCX
def generate_office_admin_docx(doc, data, candidate_name, candidate_id, position, scores):
    render_header(doc, position, candidate_name, candidate_id, data["exam_date"])
    q_num = 1

    if len(scores) > 0 and scores[0] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Phần thi Microsoft Word: Soạn thảo văn bản hành chính ({scores[0]}đ):", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, f"   Yêu cầu: Soạn thảo văn bản '{data['doc_title']}' đúng quy định Nghị định 30/2020/NĐ-CP.", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, "   - Đảm bảo font chữ Arial/Times New Roman, căn lề đúng tiêu chuẩn, có quốc hiệu tiêu ngữ và chỗ ký.", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    if len(scores) > 1 and scores[1] > 0:
        add_paragraph_with_run(doc, f"Câu {q_num}) Phần thi Microsoft Excel: Xử lý dữ liệu bảng tính ({scores[1]}đ):", bold=True, size_pt=11, space_after=4)
        add_paragraph_with_run(doc, f"   Yêu cầu: Mở file Excel dữ liệu {data['excel_dataset_count']} khách hàng đi kèm đề thi.", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, "   - Dùng VLOOKUP để điền tên khoa từ sheet DanhMucKhoa; dùng SUM, AVERAGE và IF để tính toán, phân loại chi phí.", size_pt=11, space_after=2)
        add_paragraph_with_run(doc, "   - Trích lọc danh sách bệnh nhân có chi phí > 1.000.000đ và vẽ biểu đồ cột thể hiện chi phí.", size_pt=11, space_after=4)
        add_paragraph_with_run(doc, "Số điểm đạt được:………(đ)", size_pt=10, italic=True, space_after=12)
        q_num += 1

    add_paragraph_with_run(doc, f"Lưu ý: Nộp file bài làm Word và Excel vào thư mục BaiLam_{candidate_name}.", bold=True, size_pt=11)


def generate_office_excel_file(output_path, row_count=20):
    """Creates a synthetic Excel exercise dataset independent from HIS."""
    workbook = openpyxl.Workbook()
    data_sheet = workbook.active
    data_sheet.title = "DuLieu"
    department_sheet = workbook.create_sheet("DanhMucKhoa")

    departments = [
        ("K01", "Khoa Khám bệnh"),
        ("K02", "Khoa Nội"),
        ("K03", "Khoa Ngoại tổng hợp"),
        ("K04", "Khoa Phụ Sản"),
        ("K05", "Khoa Nhi"),
    ]
    department_sheet.append(["MaKhoa", "TenKhoa"])
    for department in departments:
        department_sheet.append(department)

    headers = [
        "STT",
        "MaHoSo",
        "HoTen",
        "MaKhoa",
        "TenKhoa",
        "ChiPhiKham",
        "TyLeBaoHiem",
        "BaoHiemChiTra",
        "PhanLoaiChiPhi",
    ]
    data_sheet.append(headers)
    insurance_rates = [0.0, 0.6, 0.8, 0.95, 1.0]
    for index in range(1, row_count + 1):
        department_code = departments[(index - 1) % len(departments)][0]
        examination_cost = 350_000 + ((index * 173_000) % 1_850_000)
        data_sheet.append(
            [
                index,
                f"HS{index:03d}",
                f"Khách hàng {index:02d}",
                department_code,
                "",
                examination_cost,
                insurance_rates[(index - 1) % len(insurance_rates)],
                "",
                "",
            ]
        )

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = ExcelFont(color="FFFFFF", bold=True)
    for sheet in (data_sheet, department_sheet):
        for cell in sheet[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = sheet.dimensions

    widths = [7, 12, 22, 10, 24, 16, 15, 18, 20]
    for index, width in enumerate(widths, start=1):
        data_sheet.column_dimensions[get_column_letter(index)].width = width
    department_sheet.column_dimensions["A"].width = 12
    department_sheet.column_dimensions["B"].width = 28

    for row in range(2, row_count + 2):
        data_sheet.cell(row, 6).number_format = '#,##0" đ"'
        data_sheet.cell(row, 7).number_format = "0%"
        data_sheet.cell(row, 8).number_format = '#,##0" đ"'

    workbook.save(output_path)
    return output_path


# MAIN ENTRY POINT FOR DOCX GENERATION
def generate_docx_file(data, template_type, candidate_name, candidate_id, position, output_path, scores=None, dept=""):
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.75)
        
    if scores is None:
        scores = []

    # Check for modular action results
    if "action_results" in data and data["action_results"]:
        from exam_actions import ACTION_REGISTRY
        render_header(doc, position, candidate_name, candidate_id, data.get("exam_date", ""), dept=dept or data.get("dept_name", ""))
        q_num = 1
        for item in data["action_results"]:
            act_code = item.get("action_code")
            act_data = item.get("data", {})
            score = item.get("score", 0)
            if score <= 0:
                continue
            
            action_spec = ACTION_REGISTRY.get(act_code)
            if action_spec and "render_docx" in action_spec:
                action_spec["render_docx"](doc, act_data, score, q_num)
                add_paragraph_with_run(doc, f"Điểm đạt được: …… / {score} đ", font_name="Times New Roman", size_pt=10.5, italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=14)
                q_num += 1
        
        doc.save(output_path)
        try:
            print(f"Generated modular docx at: {output_path}")
        except Exception:
            pass
        return output_path
        
    if template_type == "NURSE_INPATIENT":
        generate_nurse_inpatient_docx(doc, data, candidate_name, candidate_id, position, scores)
    elif template_type == "RECEPTIONIST_INPATIENT":
        generate_receptionist_inpatient_docx(doc, data, candidate_name, candidate_id, position, scores)
    elif template_type == "NURSE_DIRECT_RECEPTION":
        generate_nurse_direct_reception_docx(doc, data, candidate_name, candidate_id, position, scores)
    elif template_type == "RECEPTIONIST_DIRECT":
        generate_receptionist_direct_docx(doc, data, candidate_name, candidate_id, position, scores)
    elif template_type == "NURSE_OUTPATIENT":
        generate_nurse_outpatient_docx(doc, data, candidate_name, candidate_id, position, scores)
    elif template_type == "TECHNICIAN":
        generate_technician_docx(doc, data, candidate_name, candidate_id, position, scores)
    elif template_type == "CASHIER_COUNTER":
        generate_cashier_counter_docx(doc, data, candidate_name, candidate_id, position, scores)
    elif template_type == "OFFICE_ADMIN":
        generate_office_admin_docx(doc, data, candidate_name, candidate_id, position, scores)
    else:
        # Fallback basic template
        add_paragraph_with_run(doc, "BÀI THI VI TÍNH", bold=True, size_pt=16, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_paragraph_with_run(doc, f"Thí sinh: {candidate_name} - SBD: {candidate_id}", bold=True, size_pt=11)
        add_paragraph_with_run(doc, f"Vị trí: {position}", size_pt=11)
        
    doc.save(output_path)
    try:
        print(f"Generated docx at: {output_path}")
    except Exception:
        print("Generated docx successfully.")
    return output_path
