import random
from datetime import datetime, timedelta
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_font(run, font_name="Times New Roman", size_pt=11, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic

def add_paragraph_with_run(doc, text="", font_name="Times New Roman", size_pt=11, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4, line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        set_font(run, font_name, size_pt, bold, italic)
    return p

def get_cutoff_date(exam_date_str, days_back=30):
    try:
        exam_dt = datetime.strptime(exam_date_str, "%d/%m/%Y")
    except Exception:
        exam_dt = datetime.now()
    cutoff_dt = exam_dt - timedelta(days=days_back)
    start_dt = cutoff_dt - timedelta(days=60)
    return start_dt.strftime("%d/%m/%Y"), cutoff_dt.strftime("%d/%m/%Y")


# --- ACTION IMPLEMENTATIONS (WITH SHARED CANDIDATE CONTEXT) ---

# 1. NT_NHAN_BENH_KHOA: Nhận bệnh nhân vào khoa điều trị (Hàng chờ HIS)
def prepare_nhan_benh_khoa(cm, mm, dept_name, exam_date, params, context=None):
    if context and context.get("bn_bhyt"):
        bn_bhyt = context["bn_bhyt"]
    else:
        bn_bhyt = cm.get_patients(1, must_have_bhyt=True, valid_on=exam_date)[0]
        if context is not None:
            context["bn_bhyt"] = bn_bhyt

    if context and context.get("bn_vp"):
        bn_vp = context["bn_vp"]
    else:
        bn_vp = cm.get_patients(1, must_have_bhyt=False)[0]
        if context is not None:
            context["bn_vp"] = bn_vp

    return {
        "bn_bhyt": bn_bhyt,
        "bn_vp": bn_vp,
        "dept_name": dept_name,
        "exam_date": exam_date
    }

def render_docx_nhan_benh_khoa(doc, data, score, q_index):
    add_paragraph_with_run(doc, f"Câu {q_index}) Nhận bệnh nhân vào khoa điều trị ({score} điểm):", bold=True, size_pt=11.5, space_after=4)
    
    bn_bhyt = data["bn_bhyt"]
    birth_bhyt = bn_bhyt.get("NgaySinh") or bn_bhyt.get("NamSinh") or ""
    doi_tuong_bhyt = bn_bhyt.get("DoiTuong", "BHYT")
    dkkcb_bhyt = bn_bhyt.get("DKKCB") or "66232"
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(2)
    p1.paragraph_format.left_indent = Inches(0.2)
    p1.paragraph_format.line_spacing = 1.15
    r1 = p1.add_run("a. Bệnh nhân BHYT: ")
    set_font(r1, font_name="Times New Roman", bold=True, size_pt=11)
    r1_text = p1.add_run(f"Họ tên: {bn_bhyt['TenBenhNhan']}  -  Ngày sinh: {birth_bhyt}  -  Giới tính: {bn_bhyt.get('GioiTinh', 'Nam')}\n   Số thẻ BHYT: {bn_bhyt.get('SoBHYT', '')} (Nơi ĐKKCB: {dkkcb_bhyt})  -  Đối tượng: {doi_tuong_bhyt}")
    set_font(r1_text, font_name="Times New Roman", size_pt=11)
    
    bn_vp = data["bn_vp"]
    birth_vp = bn_vp.get("NgaySinh") or bn_vp.get("NamSinh") or ""
    doi_tuong_vp = bn_vp.get("DoiTuong", "Viện phí")
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run("b. Bệnh nhân Viện phí: ")
    set_font(r2, font_name="Times New Roman", bold=True, size_pt=11)
    r2_text = p2.add_run(f"Họ tên: {bn_vp['TenBenhNhan']}  -  Ngày sinh: {birth_vp}  -  Giới tính: {bn_vp.get('GioiTinh', 'Nam')}  -  Đối tượng: {doi_tuong_vp}")
    set_font(r2_text, font_name="Times New Roman", size_pt=11)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(4)
    p3.paragraph_format.left_indent = Inches(0.2)
    p3.paragraph_format.line_spacing = 1.2
    r3 = p3.add_run("Yêu cầu thí sinh thực hiện nhận bệnh nhân vào khoa trên phần mềm HIS và ghi nhận kết quả:\n")
    set_font(r3, font_name="Times New Roman", italic=True, size_pt=10.5)
    r3_svv1 = p3.add_run("- Số vào viện (SVV) Bệnh nhân BHYT       : ………………………………………………………………\n")
    set_font(r3_svv1, font_name="Times New Roman", size_pt=11)
    r3_svv2 = p3.add_run("- Số vào viện (SVV) Bệnh nhân Viện phí : ………………………………………………………………")
    set_font(r3_svv2, font_name="Times New Roman", size_pt=11)


# 2. TN_TIEP_NHAN: Tự tiếp nhận mới bệnh nhân
def prepare_tiep_nhan(cm, mm, dept_name, exam_date, params, context=None):
    if context and context.get("bn_bhyt"):
        bn_bhyt = context["bn_bhyt"]
    else:
        bn_bhyt = cm.get_patients(1, must_have_bhyt=True, valid_on=exam_date)[0]
        if context is not None:
            context["bn_bhyt"] = bn_bhyt

    if context and context.get("bn_vp"):
        bn_vp = context["bn_vp"]
    else:
        bn_vp = cm.get_patients(1, must_have_bhyt=False)[0]
        if context is not None:
            context["bn_vp"] = bn_vp

    return {
        "bn_bhyt": bn_bhyt,
        "bn_vp": bn_vp,
        "dept_name": dept_name,
        "exam_date": exam_date
    }

def render_docx_tiep_nhan(doc, data, score, q_index):
    add_paragraph_with_run(doc, f"Câu {q_index}) Trực tiếp tiếp nhận thông tin bệnh nhân ({score} điểm):", bold=True, size_pt=11.5, space_after=4)
    
    bn_bhyt = data["bn_bhyt"]
    birth_bhyt = bn_bhyt.get("NgaySinh") or bn_bhyt.get("NamSinh") or ""
    addr_bhyt = bn_bhyt.get("DiaChiLienHe") or bn_bhyt.get("DiaChi") or "Buôn Ma Thuột, Đắk Lắk"
    han_bhyt = f"{bn_bhyt.get('BHYTTuNgay', '')} - {bn_bhyt.get('BHYTDenNgay', '')}".strip(" -")
    dkkcb_bhyt = bn_bhyt.get("DKKCB") or "66232"
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(2)
    p1.paragraph_format.left_indent = Inches(0.2)
    p1.paragraph_format.line_spacing = 1.15
    r1 = p1.add_run("a. Bệnh nhân BHYT:\n")
    set_font(r1, font_name="Times New Roman", bold=True, size_pt=11)
    r1_text = p1.add_run(f"- Họ và tên: {bn_bhyt['TenBenhNhan']}  -  Ngày sinh: {birth_bhyt} ({bn_bhyt.get('GioiTinh', 'Nam')})\n- Số thẻ BHYT: {bn_bhyt.get('SoBHYT', '')} (Hạn thẻ: {han_bhyt}  -  Nơi ĐKKCB: {dkkcb_bhyt})\n- Địa chỉ: {addr_bhyt}\n- Lý do vào viện / Chẩn đoán: {bn_bhyt.get('LyDoKham', 'Đau bụng cấp / Theo dõi viêm ruột thừa')}")
    set_font(r1_text, font_name="Times New Roman", size_pt=11)
    
    bn_vp = data["bn_vp"]
    birth_vp = bn_vp.get("NgaySinh") or bn_vp.get("NamSinh") or ""
    addr_vp = bn_vp.get("DiaChiLienHe") or bn_vp.get("DiaChi") or "Buôn Ma Thuột, Đắk Lắk"
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run("b. Bệnh nhân Viện phí:\n")
    set_font(r2, font_name="Times New Roman", bold=True, size_pt=11)
    r2_text = p2.add_run(f"- Họ và tên: {bn_vp['TenBenhNhan']}  -  Ngày sinh: {birth_vp} ({bn_vp.get('GioiTinh', 'Nam')})\n- Địa chỉ: {addr_vp}\n- Lý do vào viện / Chẩn đoán: {bn_vp.get('LyDoKham', 'Sốt cao, ho kéo dài')}")
    set_font(r2_text, font_name="Times New Roman", size_pt=11)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(4)
    p3.paragraph_format.left_indent = Inches(0.2)
    p3.paragraph_format.line_spacing = 1.2
    r3 = p3.add_run("Yêu cầu thí sinh nhập mới thông tin bệnh nhân trên phần mềm HIS và ghi lại:\n")
    set_font(r3, font_name="Times New Roman", italic=True, size_pt=10.5)
    r3_svv1 = p3.add_run("- Mã y tế / Số tiếp nhận Bn BHYT       : ………………………………………………………………\n")
    set_font(r3_svv1, font_name="Times New Roman", size_pt=11)
    r3_svv2 = p3.add_run("- Mã y tế / Số tiếp nhận Bn Viện phí : ………………………………………………………………")
    set_font(r3_svv2, font_name="Times New Roman", size_pt=11)


# 3. YL_CHI_DINH_CLS: Chỉ định Cận lâm sàng (Thực hiện trên Bệnh nhân BHYT)
def prepare_chi_dinh_cls(cm, mm, dept_name, exam_date, params, context=None):
    bn_bhyt = context.get("bn_bhyt") if context else None
    if not bn_bhyt:
        bn_bhyt = cm.get_patients(1, must_have_bhyt=True, valid_on=exam_date)[0]
        if context is not None:
            context["bn_bhyt"] = bn_bhyt

    allowed_groups = mm.get_services_for_dept(dept_name)
    counts_config = params.get("counts_by_group", {})
    selected_services = []
    
    if counts_config:
        for grp, count in counts_config.items():
            svcs = cm.get_services(int(count), nhom=grp)
            selected_services.extend(svcs)
    else:
        img_groups = [g for g in allowed_groups if any(k in g.lower() for k in ["siêu âm", "x-quang", "ct", "mri", "nội soi", "thủ thuật"])]
        lab_groups = [g for g in allowed_groups if any(k in g.lower() for k in ["xét nghiệm", "huyết học", "hóa sinh", "vi sinh", "khí máu"])]
        
        if img_groups:
            img_services = cm.get_services(1, nhom=img_groups)
            selected_services.extend(img_services)
        if lab_groups:
            lab_services = cm.get_services(3, nhom=lab_groups)
            selected_services.extend(lab_services)
        
        if not selected_services:
            selected_services = cm.get_services(3, nhom=allowed_groups if allowed_groups else None)

    if context is not None:
        context["ordered_services"] = selected_services

    return {
        "services": selected_services,
        "bn_bhyt": bn_bhyt,
        "dept_name": dept_name
    }

def render_docx_chi_dinh_cls(doc, data, score, q_index):
    bn_bhyt = data.get("bn_bhyt")
    pt_name = f" cho bệnh nhân BHYT ({bn_bhyt['TenBenhNhan']})" if bn_bhyt else ""
    add_paragraph_with_run(doc, f"Câu {q_index}) Chỉ định các dịch vụ Cận lâm sàng{pt_name} ({score} điểm):", bold=True, size_pt=11.5, space_after=4)
    services = data.get("services", [])
    for idx, s in enumerate(services, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(f"- {s['TenDichVu']} ({s.get('NhomDichVu', '')})")
        set_font(r, font_name="Times New Roman", size_pt=11)


# 4. YL_CHI_DINH_THUOC_VTYT: Kê đơn / Lên y lệnh Thuốc & VTYT (Thực hiện trên Bệnh nhân BHYT)
def prepare_chi_dinh_thuoc_vtyt(cm, mm, dept_name, exam_date, params, context=None):
    bn_bhyt = context.get("bn_bhyt") if context else None
    if not bn_bhyt:
        bn_bhyt = cm.get_patients(1, must_have_bhyt=True, valid_on=exam_date)[0]
        if context is not None:
            context["bn_bhyt"] = bn_bhyt

    warehouses = mm.get_pharmacies_for_dept(dept_name)
    num_drugs = int(params.get("num_drugs", 3))
    num_supplies = int(params.get("num_supplies", 2))
    
    drugs = cm.get_drugs(num_drugs, nguon="BH", min_ton=5.0, warehouses=warehouses)
    supplies = cm.get_drugs(num_supplies, nguon="VP", min_ton=5.0, warehouses=warehouses)

    if context is not None:
        context["ordered_drugs"] = drugs
        context["ordered_supplies"] = supplies

    return {
        "drugs": drugs,
        "supplies": supplies,
        "bn_bhyt": bn_bhyt,
        "dept_name": dept_name
    }

def render_docx_chi_dinh_thuoc_vtyt(doc, data, score, q_index):
    bn_bhyt = data.get("bn_bhyt")
    pt_name = f" cho bệnh nhân BHYT ({bn_bhyt['TenBenhNhan']})" if bn_bhyt else ""
    add_paragraph_with_run(doc, f"Câu {q_index}) Lên y lệnh Thuốc và Vật tư y tế{pt_name} ({score} điểm):", bold=True, size_pt=11.5, space_after=4)
    
    drugs = data.get("drugs", [])
    if drugs:
        p_d = doc.add_paragraph()
        p_d.paragraph_format.space_after = Pt(2)
        p_d.paragraph_format.left_indent = Inches(0.2)
        p_d.paragraph_format.line_spacing = 1.15
        r_d = p_d.add_run("+ Thuốc Bảo hiểm y tế:")
        set_font(r_d, font_name="Times New Roman", bold=True, size_pt=11)
        for d in drugs:
            dvt = d.get('DVTTinh') or 'Viên'
            sl = random.randint(1, 3)
            cachdung = f"Sáng 1 {dvt}, chiều 1 {dvt}." if sl >= 2 else f"Sáng 1 {dvt}."
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(f"- {d['TenDuoc']} ({dvt}) - Số lượng: {sl} - Cách dùng: {cachdung}")
            set_font(r, font_name="Times New Roman", size_pt=11)
            
    supplies = data.get("supplies", [])
    if supplies:
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_after = Pt(2)
        p_s.paragraph_format.left_indent = Inches(0.2)
        p_s.paragraph_format.line_spacing = 1.15
        r_s = p_s.add_run("+ Vật tư y tế / Thuốc Viện phí:")
        set_font(r_s, font_name="Times New Roman", bold=True, size_pt=11)
        for s in supplies:
            svt = s.get('DVTTinh') or 'Cái'
            sl = random.randint(1, 2)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(f"- {s['TenDuoc']} ({svt}) - Số lượng: {sl}")
            set_font(r, font_name="Times New Roman", size_pt=11)


# 5. YL_TRA_THUOC: Trả thuốc thừa / Hủy y lệnh (Thực hiện trên Thuốc đã kê của Bệnh nhân BHYT)
def prepare_tra_thuoc(cm, mm, dept_name, exam_date, params, context=None):
    bn_bhyt = context.get("bn_bhyt") if context else None
    ordered_drugs = context.get("ordered_drugs", []) if context else []
    ordered_supplies = context.get("ordered_supplies", []) if context else []
    
    if ordered_drugs:
        drug = random.choice(ordered_drugs)
    elif ordered_supplies:
        drug = random.choice(ordered_supplies)
    else:
        warehouses = mm.get_pharmacies_for_dept(dept_name)
        drug = cm.get_drugs(1, nguon="BH", warehouses=warehouses)[0]

    return {
        "drug": drug,
        "quantity": 1,
        "bn_bhyt": bn_bhyt,
        "dept_name": dept_name
    }

def render_docx_tra_thuoc(doc, data, score, q_index):
    bn_bhyt = data.get("bn_bhyt")
    pt_name = f" cho bệnh nhân BHYT ({bn_bhyt['TenBenhNhan']})" if bn_bhyt else ""
    add_paragraph_with_run(doc, f"Câu {q_index}) Thực hiện trả thuốc / Hủy y lệnh{pt_name} ({score} điểm):", bold=True, size_pt=11.5, space_after=4)
    d = data["drug"]
    dvt = d.get('DVTTinh') or 'Viên'
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"- Hoàn trả lại tủ trực / kho: {d['TenDuoc']} ({dvt}) - Số lượng hoàn trả: {data['quantity']} {dvt} (Lý do: Bệnh nhân đỡ đau / Đổi phác đồ điều trị).")
    set_font(r, font_name="Times New Roman", size_pt=11)


# 6. YL_DOI_THEM_DICH_VU: Đổi hoặc bổ sung dịch vụ CLS (Đổi từ dịch vụ đã chỉ định trước đó)
def prepare_doi_them_dich_vu(cm, mm, dept_name, exam_date, params, context=None):
    bn_bhyt = context.get("bn_bhyt") if context else None
    ordered = context.get("ordered_services", []) if context else []
    
    # 1. Choose swap_out from already ordered services
    if ordered:
        swap_out = random.choice(ordered)
    else:
        allowed_groups = mm.get_services_for_dept(dept_name)
        svcs = cm.get_services(1, nhom=allowed_groups if allowed_groups else None)
        swap_out = svcs[0] if svcs else cm.get_services(1)[0]

    # 2. Choose swap_in and added from allowed groups not in ordered
    allowed_groups = mm.get_services_for_dept(dept_name)
    ordered_ids = {s.get("MaDichVu") for s in ordered} | {swap_out.get("MaDichVu")}
    
    available_svcs = [
        s for s in cm.services 
        if s.get("MaDichVu") not in ordered_ids and (not allowed_groups or s.get("NhomDichVu") in allowed_groups)
    ]
    
    if len(available_svcs) >= 2:
        sampled = random.sample(available_svcs, 2)
        swap_in, added = sampled[0], sampled[1]
    elif len(available_svcs) == 1:
        swap_in = available_svcs[0]
        added = cm.get_services(1)[0]
    else:
        swap_in = cm.get_services(1)[0]
        added = cm.get_services(1)[0]

    return {
        "swap_out": swap_out,
        "swap_in": swap_in,
        "added": added,
        "bn_bhyt": bn_bhyt,
        "dept_name": dept_name
    }

def render_docx_doi_them_dich_vu(doc, data, score, q_index):
    bn_bhyt = data.get("bn_bhyt")
    pt_name = f" cho bệnh nhân BHYT ({bn_bhyt['TenBenhNhan']})" if bn_bhyt else ""
    add_paragraph_with_run(doc, f"Câu {q_index}) Thay đổi và bổ sung dịch vụ kỹ thuật{pt_name} ({score} điểm):", bold=True, size_pt=11.5, space_after=4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f"- Hủy chỉ định dịch vụ: {data['swap_out']['TenDichVu']}\n")
    r2 = p.add_run(f"- Đổi sang chỉ định dịch vụ: {data['swap_in']['TenDichVu']}\n")
    r3 = p.add_run(f"- Chỉ định bổ sung thêm: {data['added']['TenDichVu']}")
    set_font(r1, font_name="Times New Roman", size_pt=11)
    set_font(r2, font_name="Times New Roman", size_pt=11)
    set_font(r3, font_name="Times New Roman", size_pt=11)


# 7. TK_KIEM_TON_KHO: Tra cứu tồn kho Dược / Tủ trực
def prepare_kiem_ton_kho(cm, mm, dept_name, exam_date, params, context=None):
    warehouses = mm.get_pharmacies_for_dept(dept_name)
    drug = cm.get_drugs(1, warehouses=warehouses)[0]
    return {
        "drug": drug
    }

def render_docx_kiem_ton_kho(doc, data, score, q_index):
    add_paragraph_with_run(doc, f"Câu {q_index}) Tra cứu và kiểm tra số lượng tồn kho tủ trực ({score} điểm):", bold=True, size_pt=11.5, space_after=4)
    d = data["drug"]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"- Kiểm tra số lượng tồn khả dụng hiện tại của mặt hàng: {d['TenDuoc']} (Mã: {d['MaDuoc']}) trong kho / tủ trực khoa.\n")
    set_font(r, font_name="Times New Roman", size_pt=11)
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.paragraph_format.line_spacing = 1.2
    r_blank1 = p2.add_run("- Số lượng tồn thực tế tra cứu được : ………………………………………………………………\n")
    set_font(r_blank1, font_name="Times New Roman", size_pt=11)
    r_blank2 = p2.add_run("- Ghi chú / Tên kho kiểm tra         : ………………………………………………………………")
    set_font(r_blank2, font_name="Times New Roman", size_pt=11)


# 8. CK_CHUYEN_KHOA: Chuyển khoa điều trị ca cũ (-30 ngày)
def prepare_chuyen_khoa(cm, mm, dept_name, exam_date, params, context=None):
    start_date, end_date = get_cutoff_date(exam_date, days_back=30)
    all_depts = [
        "Khoa Ngoại tổng hợp", "Khoa Chấn thương chỉnh hình", 
        "Khoa Hồi sức tích cực", "Khoa Nội", "Khoa Phụ Sản", "Khoa Nhi"
    ]
    targets = [d for d in all_depts if d != dept_name]
    target_dept = random.choice(targets) if targets else "Khoa Hồi sức tích cực"
    return {
        "start_date": start_date,
        "end_date": end_date,
        "target_dept": target_dept
    }

def render_docx_chuyen_khoa(doc, data, score, q_index):
    add_paragraph_with_run(doc, f"Câu {q_index}) Chuyển khoa điều trị ({score} điểm):", bold=True, size_pt=11.5, space_after=4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f"- Tìm 01 bệnh nhân bất kỳ đã vào khoa trong khoảng thời gian từ ngày {data['start_date']} đến ngày {data['end_date']}.\n")
    r2 = p.add_run(f"- Thực hiện chuyển bệnh nhân sang {data['target_dept']}.\n")
    set_font(r1, font_name="Times New Roman", size_pt=11)
    set_font(r2, font_name="Times New Roman", size_pt=11)
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.paragraph_format.line_spacing = 1.2
    r_b1 = p2.add_run("- Họ và tên bệnh nhân đã chuyển  : ………………………………………………………………\n")
    r_b2 = p2.add_run("- Mã bệnh án / Số vào viện (SVV) : ………………………………………………………………")
    set_font(r_b1, font_name="Times New Roman", size_pt=11)
    set_font(r_b2, font_name="Times New Roman", size_pt=11)


# 9. RV_CHO_RA_VIEN: Cho ra viện ca cũ (-30 ngày)
def prepare_cho_ra_vien(cm, mm, dept_name, exam_date, params, context=None):
    start_date, end_date = get_cutoff_date(exam_date, days_back=30)
    return {
        "start_date": start_date,
        "end_date": end_date
    }

def render_docx_cho_ra_vien(doc, data, score, q_index):
    add_paragraph_with_run(doc, f"Câu {q_index}) Cho bệnh nhân ra viện ({score} điểm):", bold=True, size_pt=11.5, space_after=4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f"- Tìm 01 bệnh nhân đang nằm điều trị nội trú trong khoảng thời gian từ ngày {data['start_date']} đến {data['end_date']}.\n")
    r2 = p.add_run("- Thực hiện các thủ tục cho bệnh nhân xuất viện trên phần mềm HIS.\n")
    set_font(r1, font_name="Times New Roman", size_pt=11)
    set_font(r2, font_name="Times New Roman", size_pt=11)
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.paragraph_format.line_spacing = 1.2
    r_b1 = p2.add_run("- Họ và tên bệnh nhân cho ra viện : ………………………………………………………………\n")
    r_b2 = p2.add_run("- Mã bệnh án / Số vào viện (SVV) : ………………………………………………………………\n")
    r_b3 = p2.add_run("- Ngày giờ xuất viện ghi nhận    : ………………………………………………………………")
    set_font(r_b1, font_name="Times New Roman", size_pt=11)
    set_font(r_b2, font_name="Times New Roman", size_pt=11)
    set_font(r_b3, font_name="Times New Roman", size_pt=11)


# 10. TC_THU_TAM_UNG: Thu tạm ứng nội trú / ngoại trú
def prepare_thu_tam_ung(cm, mm, dept_name, exam_date, params, context=None):
    cases = [
        {"ma_ba": "26004776", "tien": "1,500,000"},
        {"ma_ba": "26001724/NG", "tien": "500,000"}
    ]
    return {
        "cases": cases
    }

def render_docx_thu_tam_ung(doc, data, score, q_index):
    add_paragraph_with_run(doc, f"Câu {q_index}) Thu tạm ứng viện phí ({score} điểm):", bold=True, size_pt=11.5, space_after=4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.line_spacing = 1.15
    r_intro = p.add_run("Tìm và lập phiếu thu tạm ứng viện phí cho các ca bệnh sau:\n")
    set_font(r_intro, font_name="Times New Roman", size_pt=11)
    for idx, c in enumerate(data["cases"], 1):
        r_c = p.add_run(f"  + Ca {idx}: Mã BN / Mã đợt khám: {c['ma_ba']}  -  Số tiền tạm ứng: {c['tien']} VNĐ\n")
        set_font(r_c, font_name="Times New Roman", size_pt=11)
        
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.paragraph_format.line_spacing = 1.2
    r_b1 = p2.add_run("- Số biên lai / Mã phiếu thu Ca 1 : ………………………………………………………………\n")
    r_b2 = p2.add_run("- Số biên lai / Mã phiếu thu Ca 2 : ………………………………………………………………")
    set_font(r_b1, font_name="Times New Roman", size_pt=11)
    set_font(r_b2, font_name="Times New Roman", size_pt=11)


# 11. TC_THANH_TOAN_RA_VIEN: Thanh toán ra viện
def prepare_thanh_toan_ra_vien(cm, mm, dept_name, exam_date, params, context=None):
    start_date, end_date = get_cutoff_date(exam_date, days_back=30)
    return {
        "start_date": start_date,
        "end_date": end_date
    }

def render_docx_thanh_toan_ra_vien(doc, data, score, q_index):
    add_paragraph_with_run(doc, f"Câu {q_index}) Thanh toán viện phí ra viện ({score} điểm):", bold=True, size_pt=11.5, space_after=4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"- Tìm 01 bệnh nhân đã có chỉ định ra viện từ ngày {data['start_date']} đến {data['end_date']}, thực hiện quyết toán chi phí và in phiếu thanh toán.\n")
    set_font(r, font_name="Times New Roman", size_pt=11)
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.paragraph_format.line_spacing = 1.2
    r_b1 = p2.add_run("- Họ và tên bệnh nhân thanh toán : ………………………………………………………………\n")
    r_b2 = p2.add_run("- Mã hóa đơn / Số phiếu thanh toán : ………………………………………………………………\n")
    r_b3 = p2.add_run("- Tổng số tiền thanh toán thực tế  : ………………………………………………………………")
    set_font(r_b1, font_name="Times New Roman", size_pt=11)
    set_font(r_b2, font_name="Times New Roman", size_pt=11)
    set_font(r_b3, font_name="Times New Roman", size_pt=11)


# 12. KQ_TRA_KET_QUA_CLS: Nhập và trả kết quả Cận lâm sàng (Huyết học / Siêu âm / Nội soi)
def prepare_tra_ket_qua_cls(cm, mm, dept_name, exam_date, params, context=None):
    sample_type = params.get("sample_type", "HUYET_HOC_18")
    patient = cm.get_patients(1)[0]
    return {
        "sample_type": sample_type,
        "patient": patient
    }

def render_docx_tra_ket_qua_cls(doc, data, score, q_index):
    add_paragraph_with_run(doc, f"Câu {q_index}) Nhập và trả kết quả Cận lâm sàng ({score} điểm):", bold=True, size_pt=11.5, space_after=4)
    pt = data["patient"]
    p_info = doc.add_paragraph()
    p_info.paragraph_format.space_after = Pt(4)
    p_info.paragraph_format.left_indent = Inches(0.2)
    p_info.paragraph_format.line_spacing = 1.15
    r_pt = p_info.add_run(f"Bệnh nhân: {pt['TenBenhNhan']} - {pt.get('NgaySinh') or pt.get('NamSinh', '')} ({pt.get('GioiTinh', 'Nam')})\n")
    set_font(r_pt, font_name="Times New Roman", bold=True, size_pt=11)
    r_msg = p_info.add_run("Nhập kết quả theo các thông số kỹ thuật bên dưới vào phần mềm HIS:")
    set_font(r_msg, font_name="Times New Roman", italic=True, size_pt=10.5)

    sample_type = data.get("sample_type", "HUYET_HOC_18")
    if sample_type == "HUYET_HOC_18":
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        headers = ["STT", "Tên chỉ số", "Kết quả", "Khoảng tham chiếu"]
        widths = [Inches(0.6), Inches(2.2), Inches(1.5), Inches(2.2)]
        
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            p_hdr = hdr_cells[i].paragraphs[0]
            p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p_hdr.runs[0], font_name="Times New Roman", bold=True, size_pt=10)
            hdr_cells[i].width = widths[i]
            set_cell_background(hdr_cells[i], "F2F2F2")
            
        params_list = [
            ("1", "WBC (Số lượng bạch cầu)", "7.2", "4.0 - 10.0 G/L"),
            ("2", "RBC (Số lượng hồng cầu)", "4.65", "3.8 - 5.5 T/L"),
            ("3", "HGB (Huyết sắc tố)", "142", "120 - 165 g/L"),
            ("4", "HCT (Hematocrit)", "0.42", "0.35 - 0.48 L/L"),
            ("5", "PLT (Số lượng tiểu cầu)", "245", "150 - 450 G/L"),
            ("6", "NEU% (Tỷ lệ Neutrophil)", "64.2", "40 - 75 %"),
            ("7", "LYM% (Tỷ lệ Lymphocyte)", "28.5", "20 - 45 %")
        ]
        for row_data in params_list:
            row_cells = table.add_row().cells
            for idx, text in enumerate(row_data):
                row_cells[idx].text = text
                row_cells[idx].width = widths[idx]
                p_cell = row_cells[idx].paragraphs[0]
                if idx in [0, 2]:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_font(p_cell.runs[0], font_name="Times New Roman", size_pt=10)
        
        p_spacer = doc.add_paragraph()
        p_spacer.paragraph_format.space_after = Pt(4)
    else:
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.space_after = Pt(4)
        p_desc.paragraph_format.left_indent = Inches(0.2)
        p_desc.paragraph_format.line_spacing = 1.15
        r_desc = p_desc.add_run("- Mô tả: Gan kích thước bình thường, bờ đều, nhu mô đồng nhất. Túi mật thành mỏng, không có sỏi. Thận hai bên không ứ nước.\n- Kết luận: Hình ảnh siêu âm ổ bụng chưa phát hiện bất thường.")
        set_font(r_desc, font_name="Times New Roman", size_pt=11)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.paragraph_format.line_spacing = 1.2
    r_b1 = p2.add_run("- Số phiếu chỉ định / Mã kết quả HIS : ………………………………………………………………\n")
    r_b2 = p2.add_run("- Trạng thái kết quả (Đã duyệt / Chưa duyệt) : ………………………………………………………………")
    set_font(r_b1, font_name="Times New Roman", size_pt=11)
    set_font(r_b2, font_name="Times New Roman", size_pt=11)


# 13. VP_SOAN_THAO_WORD: Soạn thảo văn bản hành chính theo NĐ 30
def prepare_soan_thao_word(cm, mm, dept_name, exam_date, params, context=None):
    return {}

def render_docx_soan_thao_word(doc, data, score, q_index):
    add_paragraph_with_run(doc, f"Câu {q_index}) Soạn thảo văn bản hành chính theo Nghị định 30/2020/NĐ-CP ({score} điểm):", bold=True, size_pt=11.5, space_after=4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run("- Soạn thảo một Thông báo hoặc Tờ trình nội bộ theo đúng quy chuẩn thể thức văn bản hành chính (Quốc hiệu, Tiêu ngữ, Tên cơ quan ban hành, Số/Ký hiệu, Trích yếu nội dung, Nơi nhận và Thẩm quyền ký ban hành).\n")
    set_font(r, font_name="Times New Roman", size_pt=11)
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.paragraph_format.line_spacing = 1.2
    r_b = p2.add_run("- Tên file Word đã lưu trên máy tính : ………………………………………………………………")
    set_font(r_b, font_name="Times New Roman", size_pt=11)


# 14. VP_XU_LY_EXCEL: Xử lý bảng tính Excel
def prepare_xu_ly_excel(cm, mm, dept_name, exam_date, params, context=None):
    return {}

def render_docx_xu_ly_excel(doc, data, score, q_index):
    add_paragraph_with_run(doc, f"Câu {q_index}) Xử lý dữ liệu bảng tính Excel ({score} điểm):", bold=True, size_pt=11.5, space_after=4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run("- Mở file dữ liệu Excel đi kèm trong đề thi thực hiện các yêu cầu sau:\n")
    r2 = p.add_run("  + Dùng hàm VLOOKUP để tra cứu Tên Khoa/Phòng và Đơn giá từ Sheet Danh mục sang Sheet Dữ liệu.\n")
    r3 = p.add_run("  + Dùng hàm IF / SUMIF / AVERAGE để tính Thành tiền và thống kê số liệu theo điều kiện.\n")
    r4 = p.add_run("  + Vẽ biểu đồ hình cột so sánh số liệu giữa các khoa và lưu file.\n")
    set_font(r1, font_name="Times New Roman", size_pt=11)
    set_font(r2, font_name="Times New Roman", size_pt=11)
    set_font(r3, font_name="Times New Roman", size_pt=11)
    set_font(r4, font_name="Times New Roman", size_pt=11)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.paragraph_format.line_spacing = 1.2
    r_b = p2.add_run("- Tên file Excel đã lưu trên máy tính : ………………………………………………………………")
    set_font(r_b, font_name="Times New Roman", size_pt=11)


# --- REGISTRY DICTIONARY ---

ACTION_REGISTRY = {
    "NT_NHAN_BENH_KHOA": {
        "code": "NT_NHAN_BENH_KHOA",
        "name": "Nhận bệnh nhân vào khoa điều trị (Hàng chờ HIS)",
        "category": "Tiếp nhận & Vào khoa",
        "description": "SQL tạo bệnh nhân BHYT/Viện phí vào hàng chờ để thí sinh bấm nhận bệnh vào khoa.",
        "default_score": 2.0,
        "uses_his": True,
        "needs_sql_inpatient": True,
        "prepare_data": prepare_nhan_benh_khoa,
        "render_docx": render_docx_nhan_benh_khoa
    },
    "TN_TIEP_NHAN": {
        "code": "TN_TIEP_NHAN",
        "name": "Tự tiếp nhận mới thông tin bệnh nhân",
        "category": "Tiếp nhận & Vào khoa",
        "description": "Đề in đủ thông tin 1 ca BHYT và 1 ca Viện phí để thí sinh tự nhập tay vào HIS. SQL chỉ validate.",
        "default_score": 2.5,
        "uses_his": True,
        "needs_sql_inpatient": False,
        "prepare_data": prepare_tiep_nhan,
        "render_docx": render_docx_tiep_nhan
    },
    "YL_CHI_DINH_CLS": {
        "code": "YL_CHI_DINH_CLS",
        "name": "Chỉ định dịch vụ Cận lâm sàng",
        "category": "Y lệnh Dịch vụ & Thuốc",
        "description": "Chỉ định các dịch vụ CLS theo đúng Nhóm Dịch vụ được phân quyền cho Khoa (Thực hiện trên Bệnh nhân BHYT).",
        "default_score": 2.0,
        "uses_his": True,
        "prepare_data": prepare_chi_dinh_cls,
        "render_docx": render_docx_chi_dinh_cls
    },
    "YL_CHI_DINH_THUOC_VTYT": {
        "code": "YL_CHI_DINH_THUOC_VTYT",
        "name": "Lên y lệnh Thuốc và Vật tư y tế",
        "category": "Y lệnh Dịch vụ & Thuốc",
        "description": "Lên y lệnh Thuốc BHYT và VTYT Viện phí từ kho dược tương ứng với Khoa (Thực hiện trên Bệnh nhân BHYT).",
        "default_score": 2.0,
        "uses_his": True,
        "prepare_data": prepare_chi_dinh_thuoc_vtyt,
        "render_docx": render_docx_chi_dinh_thuoc_vtyt
    },
    "YL_TRA_THUOC": {
        "code": "YL_TRA_THUOC",
        "name": "Trả thuốc thừa / Hủy y lệnh",
        "category": "Y lệnh Dịch vụ & Thuốc",
        "description": "Thực hiện hoàn trả thuốc đã được kê ở câu Lên y lệnh của Bệnh nhân BHYT.",
        "default_score": 1.0,
        "uses_his": True,
        "prepare_data": prepare_tra_thuoc,
        "render_docx": render_docx_tra_thuoc
    },
    "YL_DOI_THEM_DICH_VU": {
        "code": "YL_DOI_THEM_DICH_VU",
        "name": "Thay đổi và bổ sung dịch vụ kỹ thuật",
        "category": "Y lệnh Dịch vụ & Thuốc",
        "description": "Hủy dịch vụ đã chỉ định ở câu CLS, đổi sang dịch vụ mới và chỉ định bổ sung thêm dịch vụ.",
        "default_score": 1.0,
        "uses_his": True,
        "prepare_data": prepare_doi_them_dich_vu,
        "render_docx": render_docx_doi_them_dich_vu
    },
    "TK_KIEM_TON_KHO": {
        "code": "TK_KIEM_TON_KHO",
        "name": "Kiểm tra tồn kho Dược / Tủ trực",
        "category": "Tra cứu & Quản lý",
        "description": "Tra cứu số lượng tồn khả dụng của mặt hàng thuốc/VTYT trong tủ trực khoa.",
        "default_score": 1.0,
        "uses_his": True,
        "prepare_data": prepare_kiem_ton_kho,
        "render_docx": render_docx_kiem_ton_kho
    },
    "CK_CHUYEN_KHOA": {
        "code": "CK_CHUYEN_KHOA",
        "name": "Chuyển khoa điều trị ca cũ (-30 ngày)",
        "category": "Quản lý Người bệnh",
        "description": "Tìm bệnh nhân cũ vào khoa trước 30 ngày và thực hiện chuyển khoa.",
        "default_score": 1.0,
        "uses_his": True,
        "prepare_data": prepare_chuyen_khoa,
        "render_docx": render_docx_chuyen_khoa
    },
    "RV_CHO_RA_VIEN": {
        "code": "RV_CHO_RA_VIEN",
        "name": "Cho ra viện ca cũ (-30 ngày)",
        "category": "Quản lý Người bệnh",
        "description": "Tìm bệnh nhân cũ và thực hiện các thủ tục cho ra viện trên phần mềm.",
        "default_score": 1.0,
        "uses_his": True,
        "prepare_data": prepare_cho_ra_vien,
        "render_docx": render_docx_cho_ra_vien
    },
    "TC_THU_TAM_UNG": {
        "code": "TC_THU_TAM_UNG",
        "name": "Thu tạm ứng viện phí",
        "category": "Tài chính & Viện phí",
        "description": "Thu tiền tạm ứng nội trú / ngoại trú cho các đợt khám chỉ định.",
        "default_score": 3.0,
        "uses_his": True,
        "prepare_data": prepare_thu_tam_ung,
        "render_docx": render_docx_thu_tam_ung
    },
    "TC_THANH_TOAN_RA_VIEN": {
        "code": "TC_THANH_TOAN_RA_VIEN",
        "name": "Thanh toán viện phí ra viện",
        "category": "Tài chính & Viện phí",
        "description": "Quyết toán viện phí và in bảng kê chi phí thanh toán ra viện.",
        "default_score": 3.0,
        "uses_his": True,
        "prepare_data": prepare_thanh_toan_ra_vien,
        "render_docx": render_docx_thanh_toan_ra_vien
    },
    "KQ_TRA_KET_QUA_CLS": {
        "code": "KQ_TRA_KET_QUA_CLS",
        "name": "Nhập và trả kết quả Cận lâm sàng",
        "category": "Kỹ thuật viên & Cận lâm sàng",
        "description": "Nhập trả kết quả xét nghiệm huyết học 18 thông số hoặc mô tả siêu âm/nội soi.",
        "default_score": 6.0,
        "uses_his": True,
        "prepare_data": prepare_tra_ket_qua_cls,
        "render_docx": render_docx_tra_ket_qua_cls
    },
    "VP_SOAN_THAO_WORD": {
        "code": "VP_SOAN_THAO_WORD",
        "name": "Soạn thảo văn bản Word theo NĐ 30/2020",
        "category": "Tin học Văn phòng",
        "description": "Soạn thảo văn bản hành chính đúng chuẩn thể thức (Quốc hiệu, tiêu ngữ, ký hiệu...).",
        "default_score": 5.0,
        "uses_his": False,
        "prepare_data": prepare_soan_thao_word,
        "render_docx": render_docx_soan_thao_word
    },
    "VP_XU_LY_EXCEL": {
        "code": "VP_XU_LY_EXCEL",
        "name": "Xử lý dữ liệu bảng tính Excel",
        "category": "Tin học Văn phòng",
        "description": "Thực hành các hàm VLOOKUP, IF, SUMIF và vẽ biểu đồ trên file Excel dữ liệu giả lập.",
        "default_score": 5.0,
        "uses_his": False,
        "prepare_data": prepare_xu_ly_excel,
        "render_docx": render_docx_xu_ly_excel
    }
}

def get_action(code):
    return ACTION_REGISTRY.get(code)

def get_all_actions():
    return [
        {
            "code": a["code"],
            "name": a["name"],
            "category": a["category"],
            "description": a["description"],
            "default_score": a["default_score"],
            "uses_his": a.get("uses_his", True)
        }
        for a in ACTION_REGISTRY.values()
    ]
